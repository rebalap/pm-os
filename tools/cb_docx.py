"""
cb_docx.py — Cognitivebotics branded Word document library.

QUICK START
-----------
    from cb_docx import CB_Doc

    # From a markdown file (most common)
    CB_Doc.build_from_markdown(
        md_path  = "my-research-brief.md",
        out_path = "my-research-brief.docx",
        title    = "Research Brief: Center Lifecycle",
        date     = "April 2026",
    )

    # Programmatic — build manually
    d = CB_Doc("Interview Script: Center Owner", date="April 2026")
    d.h1("Interview Script: Center Owner / Admin")
    d.cover_table([("Product", "Autism Therapy Platform"), ("Stage", "Discovery")])
    d.h2("Section 1 — Warm-Up")
    d.body("Tell me about your center.")
    d.save("output.docx")

CLI
---
    python3 cb_docx.py input.md output.docx
    python3 cb_docx.py input.md output.docx --title "My Doc" --date "May 2026"
"""

from __future__ import annotations
import re, os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Import brand constants (same directory)
_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _dir)
import cb_brand as B


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _rgb(h: str) -> RGBColor:
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def _spacing(para, before: int = 0, after: int = 4):
    pPr = para._p.get_or_add_pPr()
    for s in pPr.findall(qn('w:spacing')):
        pPr.remove(s)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before * 20))
    sp.set(qn('w:after'),  str(after  * 20))
    pPr.append(sp)

def _indent(para, left: int = 0, hanging: int = 0):
    pPr = para._p.get_or_add_pPr()
    for i in pPr.findall(qn('w:ind')):
        pPr.remove(i)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(left))
    if hanging:
        ind.set(qn('w:hanging'), str(hanging))
    pPr.append(ind)

def _para_border(para, side: str, color: str, sz: str = '12'):
    pPr = para._p.get_or_add_pPr()
    for b in pPr.findall(qn('w:pBdr')):
        pPr.remove(b)
    pBdr = OxmlElement('w:pBdr')
    el = OxmlElement(f'w:{side}')
    el.set(qn('w:val'),   'single')
    el.set(qn('w:sz'),    sz)
    el.set(qn('w:space'), '4')
    el.set(qn('w:color'), color)
    pBdr.append(el)
    pPr.append(pBdr)

def _cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')):
        tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def _cell_borders(cell, left_color: str = None, left_sz: str = '32',
                  hide_others: bool = True):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for b in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(b)
    tcB = OxmlElement('w:tcBorders')
    sides = {'left': (left_color, left_sz)} if left_color else {}
    for side in ['top', 'right', 'bottom', 'left']:
        el = OxmlElement(f'w:{side}')
        if side in sides:
            el.set(qn('w:val'),   'single')
            el.set(qn('w:sz'),    sides[side][1])
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), sides[side][0])
        else:
            el.set(qn('w:val'),   'none' if hide_others else 'single')
            el.set(qn('w:sz'),    '0')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), 'auto')
        tcB.append(el)
    tcPr.append(tcB)

def _table_borders(table, color: str = B.BORDER, sz: str = '4'):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
    for b in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(b)
    tblB = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tblB.append(el)
    tblPr.append(tblB)

def _page_field(para):
    """Add a PAGE number field to the current paragraph."""
    run = para.add_run()
    run.font.name  = B.FONT_BODY
    run.font.size  = Pt(B.SZ_SMALL)
    run.font.color.rgb = _rgb(B.GRAY)
    for ftype, text in [('begin', None), (None, ' PAGE '), ('end', None)]:
        if ftype:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), ftype)
            run._r.append(el)
        else:
            el = OxmlElement('w:instrText')
            el.text = text
            run._r.append(el)

def _inline(para, text: str, bold: bool = False, italic: bool = False,
            font: str = B.FONT_BODY, size: int = B.SZ_BODY,
            color: str = B.BLACK):
    """Parse **bold** and *italic* markdown and add formatted runs."""
    pattern = r'(\*\*[^*]+\*\*|\*[^*\n]+\*|[^*]+)'
    for m in re.finditer(pattern, text):
        chunk = m.group(0)
        if chunk.startswith('**') and chunk.endswith('**') and len(chunk) > 4:
            inner, b, i = chunk[2:-2], True, italic
        elif chunk.startswith('*') and chunk.endswith('*') and len(chunk) > 2:
            inner, b, i = chunk[1:-1], bold, True
        else:
            inner, b, i = chunk, bold, italic
        run = para.add_run(inner)
        run.bold, run.italic = b, i
        run.font.name  = font
        run.font.size  = Pt(size)
        run.font.color.rgb = _rgb(color)


# ── CB_Doc class ──────────────────────────────────────────────────────────────

class CB_Doc:
    """
    Cognitivebotics branded Word document builder.

    Every public method returns `self` for chaining.
    Call `.save(path)` to write the file.
    """

    def __init__(self, short_title: str = "Document", date: str = "2026"):
        self._doc   = Document()
        self._title = short_title
        self._date  = date
        # Set default Normal style
        ns = self._doc.styles['Normal']
        ns.font.name  = B.FONT_BODY
        ns.font.size  = Pt(B.SZ_BODY)
        ns.font.color.rgb = _rgb(B.BLACK)
        self._setup_margins()
        self._setup_header_footer()

    # ── Page setup ────────────────────────────────────────────────────────────

    def _setup_margins(self):
        s = self._doc.sections[0]
        s.page_width   = Inches(8.5)
        s.page_height  = Inches(11)
        s.left_margin  = Inches(1.0)
        s.right_margin = Inches(1.0)
        s.top_margin   = Inches(0.9)
        s.bottom_margin = Inches(0.9)
        s.header_distance = Inches(0.4)
        s.footer_distance = Inches(0.4)

    def _setup_header_footer(self):
        sec = self._doc.sections[0]

        # Header
        hdr = sec.header
        hdr.is_linked_to_previous = False
        for p in hdr.paragraphs:
            p.clear()
        hp = hdr.paragraphs[0]
        run = hp.add_run(f"{B.PRODUCT_NAME}  |  {B.COMPANY_NAME}")
        run.bold = True
        run.font.name  = B.FONT_HEADING
        run.font.size  = Pt(B.SZ_SMALL)
        run.font.color.rgb = _rgb(B.TEAL)
        _spacing(hp, before=0, after=4)
        _para_border(hp, 'bottom', B.TEAL, sz='16')

        # Footer
        ftr = sec.footer
        ftr.is_linked_to_previous = False
        for p in ftr.paragraphs:
            p.clear()
        fp = ftr.paragraphs[0]
        _spacing(fp, before=4, after=0)
        _para_border(fp, 'top', B.BORDER, sz='4')

        def _frun(text):
            r = fp.add_run(text)
            r.font.name  = B.FONT_BODY
            r.font.size  = Pt(8)
            r.font.color.rgb = _rgb(B.GRAY)

        _frun(self._title)
        fp.add_run('\t')
        _page_field(fp)
        fp.add_run('\t')
        _frun(self._date)

        # Tab stops
        pPr = fp._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        for val, pos in [('center', '3900'), ('right', '7800')]:
            t = OxmlElement('w:tab')
            t.set(qn('w:val'), val)
            t.set(qn('w:pos'), pos)
            tabs.append(t)
        pPr.append(tabs)

    # ── Headings ──────────────────────────────────────────────────────────────

    def h1(self, text: str) -> 'CB_Doc':
        p = self._doc.add_paragraph()
        _spacing(p, before=B.SP_H1_BEFORE, after=B.SP_H1_AFTER)
        _para_border(p, 'bottom', B.TEAL, sz='16')
        run = p.add_run(text)
        run.bold = True
        run.font.name  = B.FONT_HEADING
        run.font.size  = Pt(B.SZ_H1)
        run.font.color.rgb = _rgb(B.TEAL)
        return self

    def h2(self, text: str) -> 'CB_Doc':
        p = self._doc.add_paragraph()
        _spacing(p, before=B.SP_H2_BEFORE, after=B.SP_H2_AFTER)
        run = p.add_run(text)
        run.bold = True
        run.font.name  = B.FONT_HEADING
        run.font.size  = Pt(B.SZ_H2)
        run.font.color.rgb = _rgb(B.DARK_TEAL)
        return self

    def h3(self, text: str) -> 'CB_Doc':
        p = self._doc.add_paragraph()
        _spacing(p, before=B.SP_H3_BEFORE, after=B.SP_H3_AFTER)
        run = p.add_run(text)
        run.bold = True
        run.font.name  = B.FONT_HEADING
        run.font.size  = Pt(B.SZ_H3)
        run.font.color.rgb = _rgb(B.CYAN)
        return self

    # ── Body text ─────────────────────────────────────────────────────────────

    def body(self, text: str, before: int = 0, after: int = 4) -> 'CB_Doc':
        p = self._doc.add_paragraph()
        _spacing(p, before=before, after=after)
        _inline(p, text)
        return self

    def italic(self, text: str) -> 'CB_Doc':
        """Indented italic line — for goals, probes, listen-fors."""
        p = self._doc.add_paragraph()
        _spacing(p, before=2, after=2)
        _indent(p, left=680)
        _inline(p, text, italic=True, color=B.GRAY, size=B.SZ_SMALL)
        return self

    def kv(self, key: str, value: str) -> 'CB_Doc':
        """Bold teal key + body value on same line: **Key:** value."""
        p = self._doc.add_paragraph()
        _spacing(p, before=2, after=3)
        run = p.add_run(key + ':  ')
        run.bold = True
        run.font.name  = B.FONT_HEADING
        run.font.size  = Pt(B.SZ_BODY)
        run.font.color.rgb = _rgb(B.TEAL)
        _inline(p, value)
        return self

    def spacer(self, after: int = 6) -> 'CB_Doc':
        p = self._doc.add_paragraph()
        _spacing(p, before=0, after=after)
        return self

    # ── Lists ─────────────────────────────────────────────────────────────────

    def bullet(self, text: str, level: int = 0) -> 'CB_Doc':
        p = self._doc.add_paragraph()
        _spacing(p, before=1, after=3)
        left = 360 + level * 320
        _indent(p, left=left, hanging=280)
        sym       = '●  ' if level == 0 else '–  '
        sym_color = B.CYAN if level == 0 else B.TEAL
        run = p.add_run(sym)
        run.font.name  = B.FONT_BODY
        run.font.size  = Pt(B.SZ_BODY)
        run.font.color.rgb = _rgb(sym_color)
        _inline(p, text)
        return self

    def checkbox(self, text: str) -> 'CB_Doc':
        p = self._doc.add_paragraph()
        _spacing(p, before=1, after=3)
        _indent(p, left=360, hanging=280)
        run = p.add_run('☐  ')
        run.font.name  = B.FONT_BODY
        run.font.size  = Pt(B.SZ_BODY)
        run.font.color.rgb = _rgb(B.TEAL)
        _inline(p, text)
        return self

    def question(self, number: int | str, text: str) -> 'CB_Doc':
        """Numbered interview question with teal number."""
        p = self._doc.add_paragraph()
        _spacing(p, before=10, after=3)
        _indent(p, left=360, hanging=360)
        run = p.add_run(f'{number}.   ')
        run.bold = True
        run.font.name  = B.FONT_HEADING
        run.font.size  = Pt(B.SZ_BODY)
        run.font.color.rgb = _rgb(B.TEAL)
        _inline(p, text)
        return self

    # ── Callout elements ──────────────────────────────────────────────────────

    def blockquote(self, text: str) -> 'CB_Doc':
        """Mint callout box with thick cyan left border (opening statements, quotes)."""
        tbl  = self._doc.add_table(rows=1, cols=1)
        cell = tbl.cell(0, 0)
        _cell_bg(cell, B.MINT)
        _cell_borders(cell, left_color=B.CYAN, left_sz='32')
        p = cell.paragraphs[0]
        _spacing(p, before=4, after=4)
        _indent(p, left=100)
        parts = text.split('\n')
        _inline(p, parts[0].strip(), italic=True, color=B.BLACK)
        for part in parts[1:]:
            if part.strip():
                p2 = cell.add_paragraph()
                _spacing(p2, before=4, after=4)
                _indent(p2, left=100)
                _inline(p2, part.strip(), italic=True, color=B.BLACK)
        self.spacer(6)
        return self

    def callout(self, lines: list[str], title: str = None) -> 'CB_Doc':
        """Mint callout box — for researcher reminders, notes, assumptions."""
        tbl  = self._doc.add_table(rows=1, cols=1)
        cell = tbl.cell(0, 0)
        _cell_bg(cell, B.MINT)
        _cell_borders(cell, left_color=B.CYAN, left_sz='32')
        first = cell.paragraphs[0]
        _spacing(first, before=4, after=3)
        _indent(first, left=100)
        if title:
            r = first.add_run(title + '  ')
            r.bold = True
            r.font.name  = B.FONT_HEADING
            r.font.size  = Pt(B.SZ_BODY)
            r.font.color.rgb = _rgb(B.TEAL)
        if lines:
            _inline(first, lines[0], color=B.BLACK)
            for line in lines[1:]:
                p2 = cell.add_paragraph()
                _spacing(p2, before=2, after=2)
                _indent(p2, left=100)
                _inline(p2, line, color=B.BLACK)
        self.spacer(6)
        return self

    # ── Tables ────────────────────────────────────────────────────────────────

    def table(self, headers: list[str], rows: list[list[str]],
              col_widths: list[float] = None) -> 'CB_Doc':
        """Branded table: teal header row, alternating mint/white body rows."""
        ncols = len(headers)
        tbl   = self._doc.add_table(rows=1 + len(rows), cols=ncols)
        # Header
        hrow = tbl.rows[0]
        for i, h in enumerate(headers):
            cell = hrow.cells[i]
            _cell_bg(cell, B.TEAL)
            p = cell.paragraphs[0]
            _spacing(p, before=4, after=4)
            run = p.add_run(h)
            run.bold = True
            run.font.name  = B.FONT_HEADING
            run.font.size  = Pt(B.SZ_SMALL + 1)
            run.font.color.rgb = _rgb(B.WHITE)
        # Data rows
        for r_idx, row in enumerate(rows):
            bg = B.WHITE if r_idx % 2 == 0 else B.MINT
            for c_idx, val in enumerate(row[:ncols]):
                cell = tbl.rows[r_idx + 1].cells[c_idx]
                _cell_bg(cell, bg)
                p = cell.paragraphs[0]
                _spacing(p, before=3, after=3)
                _inline(p, str(val))
        # Column widths
        if col_widths:
            for r in tbl.rows:
                for i, w in enumerate(col_widths):
                    if i < len(r.cells):
                        r.cells[i].width = Inches(w)
        _table_borders(tbl)
        self.spacer(4)
        return self

    def cover_table(self, pairs: list[tuple[str, str]]) -> 'CB_Doc':
        """Two-column metadata table for document cover blocks."""
        tbl = self._doc.add_table(rows=len(pairs), cols=2)
        for i, (k, v) in enumerate(pairs):
            bg = B.MINT if i % 2 == 0 else B.WHITE
            kc, vc = tbl.cell(i, 0), tbl.cell(i, 1)
            _cell_bg(kc, bg)
            _cell_bg(vc, bg)
            kp, vp = kc.paragraphs[0], vc.paragraphs[0]
            _spacing(kp, before=3, after=3)
            _spacing(vp, before=3, after=3)
            rk = kp.add_run(k)
            rk.bold = True
            rk.font.name  = B.FONT_HEADING
            rk.font.size  = Pt(B.SZ_SMALL + 1)
            rk.font.color.rgb = _rgb(B.TEAL)
            _inline(vp, v, font=B.FONT_BODY, size=B.SZ_SMALL + 1)
        for row in tbl.rows:
            row.cells[0].width = Inches(1.8)
            row.cells[1].width = Inches(4.7)
        _table_borders(tbl, color=B.BORDER, sz='2')
        self.spacer(8)
        return self

    # ── Save ──────────────────────────────────────────────────────────────────

    def save(self, path: str):
        self._doc.save(path)
        print(f'Saved: {path}')
        return self

    # ── Markdown builder ──────────────────────────────────────────────────────

    def from_markdown_content(self, content: str,
                               skip_top_metadata: bool = True) -> 'CB_Doc':
        """
        Parse markdown content and append styled elements to this document.

        Handles:
        - # H1  ## H2  ### H3
        - > blockquote (single or multi-paragraph)
        - | tables |
        - - [ ] checkboxes   - bullets   (indented level-2)
        - 1. numbered questions
        - indented *Probe:* / *Listen for:* → italic
        - *italic line*  **bold** inline
        - **Key:** value metadata lines
        - --- horizontal rules (ignored)
        """
        lines = content.split('\n')
        i = 0
        meta_done = False  # once True, stop skipping top metadata

        while i < len(lines):
            line     = lines[i]
            stripped = line.strip()

            if not stripped:
                i += 1
                continue

            # Horizontal rule
            if stripped == '---':
                meta_done = True
                i += 1
                continue

            # H1 — skip if it's the title already rendered in cover block
            if stripped.startswith('# ') and not stripped.startswith('## '):
                if skip_top_metadata and not meta_done:
                    i += 1
                    continue
                self.h1(stripped[2:].strip())
                i += 1
                continue

            # Skip top metadata key-value lines (e.g. **Product:** ...)
            if not meta_done and re.match(r'^\*\*[^*]+\*\*', stripped):
                i += 1
                continue

            # H2
            if stripped.startswith('## ') and not stripped.startswith('### '):
                meta_done = True
                self.h2(stripped[3:].strip())
                i += 1
                continue

            # H3
            if stripped.startswith('### '):
                meta_done = True
                self.h3(stripped[4:].strip())
                i += 1
                continue

            # Blockquote (collect all consecutive > lines)
            if stripped.startswith('>'):
                meta_done = True
                parts = []
                while i < len(lines) and lines[i].strip().startswith('>'):
                    p = lines[i].strip()
                    parts.append('' if p == '>' else p[2:])
                    i += 1
                self.blockquote('\n'.join(x for x in parts))
                continue

            # Table
            if stripped.startswith('|'):
                meta_done = True
                rows_raw = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    raw = lines[i].strip()
                    if not re.match(r'^\|[\s\-\|:]+\|$', raw):
                        cells = [c.strip() for c in raw.split('|') if c.strip()]
                        if cells:
                            rows_raw.append(cells)
                    i += 1
                if rows_raw:
                    max_c = max(len(r) for r in rows_raw)
                    rows_raw = [r + [''] * (max_c - len(r)) for r in rows_raw]
                    self.table(rows_raw[0], rows_raw[1:])
                continue

            # Checkbox
            if stripped.startswith('- [ ]'):
                meta_done = True
                self.checkbox(stripped[5:].strip())
                i += 1
                continue

            # Level-2 indented bullet
            if re.match(r'^[ \t]{3,}- ', line):
                meta_done = True
                self.bullet(re.sub(r'^[ \t]+-\s+', '', line).strip(), level=1)
                i += 1
                continue

            # Level-1 bullet
            if stripped.startswith('- '):
                meta_done = True
                self.bullet(stripped[2:])
                i += 1
                continue

            # Numbered question (1. text)
            m = re.match(r'^(\d+)\.\s+(.+)$', stripped)
            if m:
                meta_done = True
                self.question(m.group(1), m.group(2))
                i += 1
                continue

            # Indented probe / listen-for (3+ spaces then *)
            if re.match(r'^[ \t]{3,}\*', line):
                meta_done = True
                self.italic(stripped)
                i += 1
                continue

            # Standalone italic line (*text*)
            if (stripped.startswith('*') and not stripped.startswith('**')
                    and stripped.endswith('*') and not stripped.endswith('**')):
                meta_done = True
                inner = stripped[1:-1] if len(stripped) > 2 else stripped
                self.italic(inner)
                i += 1
                continue

            # Bold **Key:** value
            m_kv = re.match(r'^\*\*(.+?):\*\*\s*(.*)$', stripped)
            if m_kv:
                meta_done = True
                self.kv(m_kv.group(1), m_kv.group(2).strip())
                i += 1
                continue

            # Bold standalone **Text**
            m_b = re.match(r'^\*\*(.+)\*\*$', stripped)
            if m_b:
                meta_done = True
                p = self._doc.add_paragraph()
                _spacing(p, before=6, after=3)
                run = p.add_run(m_b.group(1))
                run.bold = True
                run.font.name  = B.FONT_HEADING
                run.font.size  = Pt(B.SZ_BODY)
                run.font.color.rgb = _rgb(B.TEAL)
                i += 1
                continue

            # Regular body text
            if stripped:
                meta_done = True
                self.body(stripped)
            i += 1

        return self

    # ── Class-level convenience constructors ─────────────────────────────────

    @classmethod
    def build_from_markdown(cls, md_path: str, out_path: str,
                             title: str = None, date: str = 'April 2026') -> str:
        """
        One-call function: read a markdown file → save a branded .docx.

        Returns the output path.

        Example
        -------
            CB_Doc.build_from_markdown(
                "research-brief.md",
                "research-brief.docx",
                title="Research Brief: Center Lifecycle",
            )
        """
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Extract title from H1 if not supplied
        if not title:
            for line in content.split('\n'):
                stripped = line.strip()
                if stripped.startswith('# ') and not stripped.startswith('## '):
                    title = stripped[2:].strip()
                    break
            title = title or os.path.splitext(os.path.basename(md_path))[0]

        doc = cls(short_title=title, date=date)

        # Extract top-level metadata (before first ---)
        md_lines = content.split('\n')
        h1_text  = ''
        meta     = []
        for line in md_lines:
            s = line.strip()
            if s.startswith('# ') and not s.startswith('## ') and not h1_text:
                h1_text = s[2:].strip()
            elif re.match(r'^\*\*(.+?):\*\*\s*(.+)$', s):
                m = re.match(r'^\*\*(.+?):\*\*\s*(.+)$', s)
                meta.append((m.group(1), m.group(2)))
            elif s == '---':
                break

        doc.h1(h1_text or title)
        if meta:
            doc.cover_table(meta)

        doc.from_markdown_content(content, skip_top_metadata=True)
        doc.save(out_path)
        return out_path


# ── CLI ───────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    import argparse
    ap = argparse.ArgumentParser(description='Convert markdown to branded .docx')
    ap.add_argument('input',          help='Input .md file')
    ap.add_argument('output',         help='Output .docx file')
    ap.add_argument('--title', '-t',  default=None, help='Document short title (footer)')
    ap.add_argument('--date',  '-d',  default='April 2026', help='Date shown in footer')
    args = ap.parse_args()
    CB_Doc.build_from_markdown(args.input, args.output,
                                title=args.title, date=args.date)
